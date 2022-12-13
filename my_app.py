from docx import Document
from docx import shared

document = Document()

document.add_picture(
    '/home/erbol/Pictures/IMG_20220927_145203_168.jpg', width=shared.Inches(2.0)
)

document.add_paragraph(
    input('What is your name? ')
)
document.add_paragraph(
    input('How old are you? ')
)
document.add_paragraph(
    input('Your phone number? ')
)

document.add_heading(
    'About me'
)
about_me = input('Tell me about yourself: ')
document.add_paragraph(about_me)

document.add_heading(
    'Work experience'
)

p = document.add_paragraph()

company = input('Which company? ')
from_date = input('From date? ')
to_date = input('To date? ')

p.add_run(company+' ').bold = True
p.add_run(from_date+'-'+to_date+'\n').italic = True

experience_detail = input(
    'Your experience in company '+company+'? '
)
p.add_run(experience_detail)

document.save('cv.docx')